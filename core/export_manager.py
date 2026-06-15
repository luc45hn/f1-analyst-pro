import io
import re
from datetime import date
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle


def _strip_markdown(text: str) -> str:
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*',     r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^>\s*',      '', text, flags=re.MULTILINE)
    text = re.sub(r'^-{3,}$',   '', text, flags=re.MULTILINE)
    return text.strip()


def _pdf_safe(text: str) -> str:
    text = re.sub(r'[^\x00-\xFF]', '', text)
    return _xml_escape(text)


def _extract_tables(text: str) -> list[dict]:
    lines = text.split("\n")
    segments = []
    i = 0
    text_buf = []

    while i < len(lines):
        line = lines[i]
        if re.match(r'^\s*\|', line) and line.rstrip().endswith('|'):
            j = i
            block = []
            while j < len(lines) and re.match(r'^\s*\|', lines[j]) and lines[j].rstrip().endswith('|'):
                block.append(lines[j])
                j += 1
            sep_idx = next(
                (k for k, l in enumerate(block)
                 if all(re.match(r'^:?-+:?$', c.strip()) for c in l.strip().strip('|').split('|') if c.strip())),
                None,
            )
            if sep_idx is not None and len(block) >= 2:
                if text_buf:
                    segments.append({"type": "text", "content": "\n".join(text_buf)})
                    text_buf = []
                rows = [
                    [c.strip() for c in l.strip().strip('|').split('|')]
                    for k, l in enumerate(block) if k != sep_idx
                ]
                segments.append({"type": "table", "rows": rows})
            else:
                text_buf.extend(block)
            i = j
        else:
            text_buf.append(line)
            i += 1

    if text_buf:
        segments.append({"type": "text", "content": "\n".join(text_buf)})
    return segments


def export_to_docx(messages: list[dict], gp_name: str, year: int) -> bytes:
    doc = Document()

    title = doc.add_heading(f"{gp_name} {year} — Análisis F1", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Exportado el {date.today().strftime('%d/%m/%Y')}")

    pending_heading = None
    for msg in messages:
        if msg["role"] == "user":
            pending_heading = msg["content"].strip()
        elif msg["role"] == "assistant":
            heading = pending_heading or "Análisis"
            doc.add_heading(heading[:120], level=1)
            pending_heading = None
            for seg in _extract_tables(msg["content"]):
                if seg["type"] == "text":
                    for line in _strip_markdown(seg["content"]).split("\n"):
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)
                else:
                    rows = seg["rows"]
                    if rows:
                        ncols = max(len(r) for r in rows)
                        tbl = doc.add_table(rows=len(rows), cols=ncols)
                        tbl.style = "Light Grid Accent 1"
                        for r_idx, row_cells in enumerate(rows):
                            for c_idx, cell_text in enumerate(row_cells):
                                cell = tbl.rows[r_idx].cells[c_idx]
                                cell.text = cell_text
                                if r_idx == 0 and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].bold = True
            chart = msg.get("chart")
            if chart is not None:
                img_bytes = chart.to_image(format="png", width=900, height=500, scale=2)
                doc.add_picture(io.BytesIO(img_bytes), width=doc.sections[0].page_width - 3600000)
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_to_pdf(messages: list[dict], gp_name: str, year: int) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2 * cm, leftMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("f1_title", parent=styles["Title"],   fontSize=16, spaceAfter=4)
    head_style  = ParagraphStyle("f1_head",  parent=styles["Heading1"], fontSize=13, spaceBefore=12, spaceAfter=6)
    body_style  = ParagraphStyle("f1_body",  parent=styles["Normal"],   fontSize=10, leading=14, spaceAfter=4)
    meta_style  = ParagraphStyle("f1_meta",  parent=styles["Normal"],   fontSize=9)

    story = [
        Paragraph(_pdf_safe(f"{gp_name} {year} — Análisis F1"), title_style),
        Paragraph(_pdf_safe(f"Exportado el {date.today().strftime('%d/%m/%Y')}"), meta_style),
        Spacer(1, 0.5 * cm),
    ]

    pending_heading = None
    for msg in messages:
        if msg["role"] == "user":
            pending_heading = msg["content"].strip()
        elif msg["role"] == "assistant":
            heading = pending_heading or "Análisis"
            story.append(Paragraph(_pdf_safe(heading[:120]), head_style))
            pending_heading = None
            for seg in _extract_tables(msg["content"]):
                if seg["type"] == "text":
                    for line in _strip_markdown(seg["content"]).split("\n"):
                        line = line.strip()
                        if line:
                            story.append(Paragraph(_pdf_safe(line), body_style))
                else:
                    rows = seg["rows"]
                    if rows:
                        pdf_data = [[_pdf_safe(c) for c in row] for row in rows]
                        tbl = Table(pdf_data, repeatRows=1)
                        tbl.setStyle(TableStyle([
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE",   (0, 0), (-1, -1), 8),
                            ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
                            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
                            ("TOPPADDING",    (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]))
                        story.append(tbl)
                        story.append(Spacer(1, 0.2 * cm))
            chart = msg.get("chart")
            if chart is not None:
                img_bytes = chart.to_image(format="png", width=900, height=500, scale=2)
                story.append(Image(io.BytesIO(img_bytes), width=15 * cm, height=8.3 * cm))
            story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    return buf.getvalue()
